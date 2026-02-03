# routes.py
from quart import Blueprint, render_template, request, redirect, url_for, flash, session
from werkzeug.security import generate_password_hash, check_password_hash
import re
from models import async_session, User, Client, Car, Part, Order, order_part
from sqlalchemy import select, func
import os
import tempfile
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import aiofiles
import aiosmtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

# Создаём Blueprint
bp = Blueprint('main', __name__)

async def get_user_by_email(email: str):
    async with async_session() as s:
        result = await s.execute(select(User).where(User.email == email))
        return result.scalar_one_or_none()

async def get_all_clients():
    async with async_session() as s:
        result = await s.execute(select(Client))
        return result.scalars().all()

async def get_client_cars(client_id: int):
    async with async_session() as s:
        result = await s.execute(select(Car).where(Car.client_id == client_id))
        return result.scalars().all()

async def get_all_parts():
    async with async_session() as s:
        result = await s.execute(select(Part))
        return result.scalars().all()

async def get_all_orders():
    async with async_session() as s:
        result = await s.execute(
            select(Order)
            .join(Client, Order.client_id == Client.id)
            .join(User, Order.user_id == User.id)
        )
        return result.scalars().all()

async def get_my_orders(user_id: int):
    async with async_session() as s:
        result = await s.execute(
            select(Order)
            .join(Client, Order.client_id == Client.id)
            .where(Order.user_id == user_id)
        )
        return result.scalars().all()

async def get_worker_orders(user_id: int):
    # В этой версии все заказы видны мастеру — можно расширить позже
    return await get_all_orders()

async def create_user(full_name: str, email: str, phone: str, password: str, role: str = 'client'):
    async with async_session() as s:
        new_user = User(
            full_name=full_name,
            email=email,
            phone=phone,
            password_hash=generate_password_hash(password),
            role=role
        )
        s.add(new_user)
        await s.commit()
        await s.refresh(new_user)
        return new_user

async def create_client(full_name: str, phone: str, email: str = None, address: str = None):
    async with async_session() as s:
        new_client = Client(
            full_name=full_name,
            phone=phone,
            email=email or None,
            address=address or None
        )
        s.add(new_client)
        await s.commit()
        await s.refresh(new_client)
        return new_client

async def create_part(name: str, price: int, stock: int):
    async with async_session() as s:
        new_part = Part(name=name.strip(), price=price, stock=stock)
        s.add(new_part)
        await s.commit()
        await s.refresh(new_part)
        return new_part

async def create_order(client_id: int, user_id: int, description: str, part_ids: list):
    async with async_session() as s:
        new_order = Order(
            client_id=client_id,
            user_id=user_id,
            description=description,
            status='new'
        )
        s.add(new_order)
        await s.flush()

        for part_id_str in part_ids:
            part_id = int(part_id_str)
            part = await s.get(Part, part_id)
            if part and part.stock > 0:
                part.stock -= 1
                await s.execute(order_part.insert().values(
                    order_id=new_order.id,
                    part_id=part_id,
                    quantity=1
                ))
            else:
                await s.rollback()
                raise ValueError(f"Недостаточно запчастей: {part.name}")
        
        await s.commit()
        return new_order

async def get_dashboard_stats():
    async with async_session() as s:
        clients_count = await s.scalar(select(func.count(Client.id)))
        users_count = await s.scalar(select(func.count(User.id)).where(User.role != 'client'))
        parts_count = await s.scalar(select(func.sum(Part.stock)))
        orders_count = await s.scalar(select(func.count(Order.id)))
        return {
            'clients_count': clients_count or 0,
            'users_count': users_count or 0,
            'parts_count': parts_count or 0,
            'orders_count': orders_count or 0,
        }
    
def generate_work_report(order_id: int, client_name: str, work_description: str, total_cost: int):
    doc = Document()
    
    # Заголовок
    title = doc.add_heading('Отчёт о выполненных работах', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Данные
    doc.add_paragraph(f"Заказ №: {order_id}")
    doc.add_paragraph(f"Клиент: {client_name}")
    doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y')}")
    doc.add_paragraph()
    
    # Описание работ
    doc.add_heading('Выполненные работы:', level=1)
    doc.add_paragraph(work_description)
    
    # Стоимость
    doc.add_paragraph()
    doc.add_paragraph(f"Итого к оплате: {total_cost} руб.", style='Intense Quote')
    
    # Подпись
    doc.add_paragraph()
    doc.add_paragraph("Мастер: _________________________")
    
    return doc

# Главная страница
@bp.route('/')
async def index():
    return await render_template('index.html')

# Регистрация
@bp.route('/register', methods=['GET', 'POST'])
async def register():
    if request.method == 'POST':
        form = await request.form
        full_name = form.get('full_name', '').strip()
        email = form.get('email', '').strip()
        phone = form.get('phone', '').strip()
        password = form.get('password', '')
        confirm_password = form.get('confirm_password', '')

        if not all([full_name, email, phone, password]):
            await flash('Все поля обязательны!', 'danger')
            return await render_template('reg.html')

        if password != confirm_password:
            await flash('Пароли не совпадают!', 'danger')
            return await render_template('reg.html')

        if not re.match(r"[^@]+@[^@]+\.[^@]+", email):
            await flash('Некорректный email!', 'danger')
            return await render_template('reg.html')

        existing = await get_user_by_email(email)
        if existing:
            await flash('Пользователь с таким email уже существует!', 'danger')
            return await render_template('reg.html')

        await create_user(full_name, email, phone, password, role='client')
        await flash('Регистрация успешна! Теперь вы можете войти.', 'success')
        return redirect(url_for('main.login'))

    return await render_template('reg.html')

# Вход для клиентов
@bp.route('/login', methods=['GET', 'POST'])
async def login():
    if request.method == 'POST':
        form = await request.form
        email = form.get('email')
        password = form.get('password')

        user = await get_user_by_email(email)
        if user and check_password_hash(user.password_hash, password):
            if user.role in ['admin', 'manager', 'master']:
                await flash('Сотрудники должны входить через специальную форму.', 'warning')
                return await render_template('login.html')
            
            session['user_id'] = user.id
            session['user_name'] = user.full_name
            session['user_role'] = user.role
            await flash(f'Добро пожаловать, {user.full_name}!', 'success')
            return redirect(url_for('main.index'))
        else:
            await flash('Неверный email или пароль', 'danger')

    return await render_template('login.html')

# Вход для сотрудников
@bp.route('/staff/login', methods=['GET', 'POST'])
async def staff_login():
    if request.method == 'POST':
        form = await request.form
        email = form.get('email')
        password = form.get('password')

        user = await get_user_by_email(email)
        if user and check_password_hash(user.password_hash, password):
            if user.role not in ['admin', 'manager', 'master']:
                await flash('Только сотрудники могут использовать эту форму.', 'danger')
                return await render_template('staff_login.html')
            
            session['user_id'] = user.id
            session['user_name'] = user.full_name
            session['user_role'] = user.role
            await flash(f'Добро пожаловать, {user.full_name}!', 'success')
            return redirect(url_for('main.index'))
        else:
            await flash('Неверный email или пароль', 'danger')

    return await render_template('staff_login.html')

@bp.route('/my_orders')
async def my_orders():
    if session.get('user_role') != 'client':
        await flash('Доступ только для клиентов.', 'danger')
        return redirect(url_for('main.index'))
    
    orders = await get_my_orders(session['user_id'])
    return await render_template('my_orders.html', orders=orders)

@bp.route('/worker_orders')
async def worker_orders():
    if session.get('user_role') not in ['master', 'admin']:
        await flash('Доступ только для мастеров.', 'danger')
        return redirect(url_for('main.index'))
    
    orders = await get_worker_orders(session['user_id'])
    return await render_template('worker_orders.html', orders=orders)

@bp.route('/warehouse')
async def warehouse():
    if session.get('user_role') not in ['admin', 'manager', 'master']:
        await flash('Доступ запрещён.', 'danger')
        return redirect(url_for('main.index'))
    
    parts = await get_all_parts()
    return await render_template('warehouse.html', parts=parts)

@bp.route('/warehouse/add', methods=['GET', 'POST'])
async def add_part():
    if session.get('user_role') not in ['admin', 'manager']:
        await flash('Только админ и менеджер могут добавлять запчасти.', 'danger')
        return redirect(url_for('main.warehouse'))

    if request.method == 'POST':
        form = await request.form
        name = form.get('name', '').strip()
        price_str = form.get('price', '').strip()
        stock_str = form.get('stock', '').strip()

        if not name or not price_str or not stock_str:
            await flash('Все поля обязательны!', 'danger')
            return await render_template('part_form.html')

        try:
            price = int(price_str)
            stock = int(stock_str)
            if price < 0 or stock < 0:
                raise ValueError
        except ValueError:
            await flash('Цена и количество должны быть целыми неотрицательными числами.', 'danger')
            return await render_template('part_form.html')

        await create_part(name, price, stock)
        await flash(f'Запчасть "{name}" добавлена!', 'success')
        return redirect(url_for('main.warehouse'))

    return await render_template('part_form.html')

@bp.route('/all_orders')
async def all_orders():
    if session.get('user_role') not in ['manager', 'admin']:
        await flash('Доступ запрещён.', 'danger')
        return redirect(url_for('main.index'))

    orders = await get_all_orders()
    return await render_template('all_orders.html', orders=orders)

@bp.route('/reports')
async def reports():
    if session.get('user_role') not in ['admin', 'manager', 'master']:
        await flash('Доступ запрещён.', 'danger')
        return redirect(url_for('main.index'))
    
    stats = await get_dashboard_stats()
    return await render_template('report.html', **stats)

# Выход
@bp.route('/logout')
async def logout():
    session.pop('user_id', None)
    session.pop('user_name', None)
    session.pop('user_role', None)
    await flash('Вы вышли из системы.', 'info')
    return redirect(url_for('main.index'))

# Профиль
@bp.route('/profil')
async def profil():
    if 'user_id' not in session:
        await flash('Пожалуйста, войдите в систему.', 'warning')
        return redirect(url_for('main.login'))
    return await render_template('profil.html')

# Список клиентов
@bp.route('/clients')
async def client_list():
    if not session.get('user_id') or session.get('user_role') not in ['admin', 'manager', 'master']:
        await flash('У вас нет доступа к этому разделу.', 'warning')
        return redirect(url_for('main.index'))

    clients = await get_all_clients()
    return await render_template('clients.html', clients=clients)

# Список пользователей (сотрудников)
@bp.route('/users')
async def user_list():
    if session.get('user_role') != 'admin':
        await flash('Доступ запрещён', 'danger')
        return redirect(url_for('main.index'))
    
    async with async_session() as s:
        result = await s.execute(select(User))
        users = result.scalars().all()
    return await render_template('users_list.html', users=users)

# Создание сотрудника
@bp.route('/users/create', methods=['GET', 'POST'])
async def create_staff():
    if session.get('user_role') != 'admin':
        await flash('Доступ запрещён', 'danger')
        return redirect(url_for('main.index'))

    if request.method == 'POST':
        form = await request.form
        full_name = form.get('full_name', '').strip()
        email = form.get('email', '').strip()
        phone = form.get('phone', '').strip()
        password = form.get('password', '')
        role = form.get('role', 'master')

        if not all([full_name, email, phone, password]):
            await flash('Все поля обязательны!', 'danger')
            return await render_template('staff_form.html')

        if not re.match(r"[^@]+@[^@]+\.[^@]+", email):
            await flash('Некорректный email!', 'danger')
            return await render_template('staff_form.html')

        existing = await get_user_by_email(email)
        if existing:
            await flash('Пользователь с таким email уже существует!', 'danger')
            return await render_template('staff_form.html')

        if role not in ['manager', 'master']:
            await flash('Недопустимая роль!', 'danger')
            return await render_template('staff_form.html')

        await create_user(full_name, email, phone, password, role=role)
        await flash(f'Сотрудник {full_name} успешно создан!', 'success')
        return redirect(url_for('main.user_list'))

    return await render_template('staff_form.html')

# Добавление клиента
@bp.route('/clients/add', methods=['GET', 'POST'])
async def add_client():
    if not session.get('user_id') or session.get('user_role') not in ['admin', 'manager', 'master']:
        await flash('У вас нет прав на добавление клиентов.', 'warning')
        return redirect(url_for('main.index'))

    if request.method == 'POST':
        form = await request.form
        full_name = form.get('full_name', '').strip()
        phone = form.get('phone', '').strip()
        email = form.get('email', '').strip()
        address = form.get('address', '').strip()

        if not full_name or not phone:
            await flash('ФИО и телефон обязательны!', 'danger')
            return await render_template('client_form.html', editing=False)

        await create_client(full_name, phone, email or None, address or None)
        await flash('Клиент успешно добавлен!', 'success')
        return redirect(url_for('main.client_list'))

    return await render_template('client_form.html', editing=False)

@bp.route('/add_order', methods=['GET', 'POST'])
async def add_order():
    if not session.get('user_id'):
        return redirect(url_for('main.login'))

    if request.method == 'POST':
        form = await request.form
        client_id = form.get('client_id')
        car_id = form.get('car_id')  # опционально
        description = form.get('description', '').strip()
        part_ids = request.form.getlist('part_ids')
        quantities = {k.replace('quantity_', ''): int(v) for k, v in form.items() if k.startswith('quantity_')}

        if not client_id:
            await flash('Выберите клиента!', 'danger')
            return redirect(url_for('main.add_order'))

        try:
            client_id = int(client_id)
        except ValueError:
            await flash('Некорректный клиент!', 'danger')
            return redirect(url_for('main.add_order'))

        # Создаём заказ с запчастями
        await create_order(client_id, session['user_id'], description, part_ids)
        await flash('Заказ создан!', 'success')
        return redirect(url_for('main.all_orders' if session.get('user_role') != 'client' else 'main.my_orders'))

    # GET: загружаем данные
    clients = await get_all_clients()
    parts = await get_all_parts()
    return await render_template('add_order.html', clients=clients, parts=parts)

@bp.route('/worker_orders/report/<int:order_id>', methods=['GET', 'POST'])
async def work_report_form(order_id):
    if session.get('user_role') not in ['master', 'admin']:
        await flash('Доступ запрещён.', 'danger')
        return redirect(url_for('main.worker_orders'))

    # Получаем заказ и клиента
    async with async_session() as s:
        order = await s.get(Order, order_id)
        if not order:
            await flash('Заказ не найден.', 'danger')
            return redirect(url_for('main.worker_orders'))
        
        client = await s.get(Client, order.client_id)
        if not client:
            await flash('Клиент не найден.', 'danger')
            return redirect(url_for('main.worker_orders'))

    if request.method == 'POST':
        form = await request.form
        work_description = form.get('work_description', '').strip()
        total_cost_str = form.get('total_cost', '').strip()
        action = form.get('action')  # 'download' или 'email'
        email_to = form.get('email_to', '').strip()

        if not work_description or not total_cost_str:
            await flash('Все поля обязательны!', 'danger')
            return await render_template('work_report_form.html', 
                                       order_id=order_id, 
                                       client=client)

        try:
            total_cost = int(total_cost_str)
            if total_cost <= 0:
                raise ValueError
        except ValueError:
            await flash('Стоимость должна быть положительным числом.', 'danger')
            return await render_template('work_report_form.html', 
                                       order_id=order_id, 
                                       client=client)

        # Генерируем документ
        doc = generate_work_report(order_id, client.full_name, work_description, total_cost)
        
        # Сохраняем во временный файл
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name

        try:
            if action == 'download':
                # Отправляем файл на скачивание
                filename = f"Отчёт_заказ_{order_id}.docx"
                return await send_file(tmp_path, as_attachment=True, attachment_filename=filename)

            elif action == 'email':
                if not email_to or not re.match(r"[^@]+@[^@]+\.[^@]+", email_to):
                    await flash('Укажите корректный email для отправки.', 'danger')
                    return await render_template('work_report_form.html', 
                                               order_id=order_id, 
                                               client=client)

                # Отправка email
                msg = MIMEMultipart()
                msg['From'] = Config.MAIL_USERNAME
                msg['To'] = email_to
                msg['Subject'] = f"Отчёт по заказу №{order_id}"

                body = f"Здравствуйте!\n\nВо вложении отчёт по выполненным работам по заказу №{order_id}."
                msg.attach(MIMEText(body, 'plain', 'utf-8'))

                # Прикрепляем файл
                with open(tmp_path, "rb") as f:
                    part = MIMEBase('application', 'vnd.openxmlformats-officedocument.wordprocessingml.document')
                    part.set_payload(f.read())
                encoders.encode_base64(part)
                part.add_header(
                    'Content-Disposition',
                    f'attachment; filename="Отчёт_заказ_{order_id}.docx"'
                )
                msg.attach(part)

                # Отправка
                await aiosmtplib.send(
                    msg,
                    hostname=Config.MAIL_SERVER,
                    port=Config.MAIL_PORT,
                    start_tls=True,
                    username=Config.MAIL_USERNAME,
                    password=Config.MAIL_PASSWORD
                )

                await flash('Отчёт успешно отправлен на email!', 'success')
                return redirect(url_for('main.worker_orders'))

        finally:
            # Удаляем временный файл
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)

    return await render_template('work_report_form.html', order_id=order_id, client=client)